from __future__ import annotations

import json
import re
from datetime import datetime
from io import BytesIO

import requests
import streamlit as st

from question_bank import (
    LESSONS,
    QUESTION_TYPES,
    TOPICS,
    TOPIC_TO_LESSON,
    assignment_payload,
    decode_payload,
    encode_payload,
    generate_questions,
    questions_to_dicts,
)


st.set_page_config(page_title="AI Quiz Chương 6 Toán 10", page_icon="AI", layout="wide")


def apply_mobile_css() -> None:
    st.markdown(
        """
        <style>
        .stApp {
            background:
                radial-gradient(circle at 12% 8%, rgba(37, 99, 235, 0.12), transparent 28%),
                radial-gradient(circle at 88% 14%, rgba(16, 185, 129, 0.12), transparent 26%),
                linear-gradient(180deg, #f8fbff 0%, #ffffff 42%, #f7fbf9 100%);
        }
        .block-container {
            max-width: 1120px;
            padding-top: 1.25rem;
        }
        .app-hero {
            border: 1px solid rgba(37, 99, 235, 0.16);
            border-radius: 18px;
            padding: 1.15rem 1.25rem;
            margin-bottom: 1rem;
            background: linear-gradient(135deg, #eff6ff 0%, #ecfdf5 55%, #fff7ed 100%);
            box-shadow: 0 12px 34px rgba(15, 23, 42, 0.07);
        }
        .app-hero h1 {
            margin: 0 0 0.35rem 0;
        }
        .app-hero p {
            color: #475569;
            margin: 0;
        }
        .color-card {
            border: 1px solid rgba(148, 163, 184, 0.35);
            border-radius: 16px;
            padding: 1rem;
            background: rgba(255,255,255,0.88);
            box-shadow: 0 10px 24px rgba(15, 23, 42, 0.055);
        }
        .question-preview {
            border-left: 6px solid #2563eb;
            background: linear-gradient(90deg, #ffffff 0%, #f8fbff 100%);
        }
        .question-card {
            border: 1px solid rgba(148, 163, 184, 0.35);
            border-left: 6px solid #2563eb;
            border-radius: 14px;
            padding: 0.95rem 1.05rem;
            margin-bottom: 0.85rem;
            background: linear-gradient(90deg, #ffffff 0%, #f8fbff 100%);
            box-shadow: 0 8px 22px rgba(15, 23, 42, 0.045);
        }
        .quiz-meta-grid {
            display: grid;
            grid-template-columns: repeat(4, minmax(0, 1fr));
            gap: 0.75rem;
            margin: 0.75rem 0 1rem;
        }
        .quiz-meta-item {
            border: 1px solid rgba(148, 163, 184, 0.28);
            border-radius: 14px;
            padding: 0.8rem;
            background: rgba(255, 255, 255, 0.84);
        }
        .quiz-meta-item span {
            display: block;
            color: #64748b;
            font-size: 0.78rem;
            margin-bottom: 0.18rem;
        }
        .quiz-meta-item strong {
            color: #0f172a;
            font-size: 1rem;
        }
        div[data-testid="stTabs"] button {
            font-weight: 700;
        }
        div[data-testid="stSidebarContent"] {
            background: linear-gradient(180deg, #eff6ff 0%, #f0fdf4 100%);
        }
        .stButton > button, .stDownloadButton > button {
            border-radius: 12px;
            font-weight: 700;
        }
        .stButton > button[kind="primary"] {
            background: linear-gradient(90deg, #2563eb 0%, #059669 100%);
            border: none;
        }
        .student-shell .stRadio [role="radiogroup"] {
            gap: 0.45rem;
        }
        .student-shell .stRadio label {
            border: 1px solid #d7dde8;
            border-radius: 10px;
            padding: 0.7rem 0.8rem;
            background: #ffffff;
            min-height: 44px;
            align-items: center;
        }
        .student-shell .stRadio label:hover {
            border-color: #2563eb;
            background: #f8fbff;
        }
        .tf-row {
            border: 1px solid #e5e7eb;
            border-radius: 10px;
            padding: 0.75rem;
            margin: 0.55rem 0;
            background: linear-gradient(90deg, #ffffff 0%, #f8fafc 100%);
        }
        .short-answer-note {
            color: #4b5563;
            font-size: 0.92rem;
            margin-bottom: 0.35rem;
        }
        .short-answer-row input {
            text-align: center;
            font-size: 1.15rem !important;
            font-weight: 700;
        }
        .student-meta {
            color: #4b5563;
            font-size: 0.95rem;
            margin-bottom: 0.75rem;
        }
        .teacher-link input {
            font-size: 0.92rem !important;
        }
        @media (max-width: 768px) {
            .block-container {
                padding: 0.8rem 0.85rem 5rem;
            }
            h1 {
                font-size: 1.55rem !important;
                line-height: 1.22 !important;
            }
            h2 {
                font-size: 1.25rem !important;
            }
            h3 {
                font-size: 1.08rem !important;
            }
            p, li, label, div[data-testid="stMarkdownContainer"] {
                font-size: 1rem;
            }
            div[data-testid="column"] {
                width: 100% !important;
                flex: 1 1 100% !important;
            }
            .student-shell .stRadio [role="radiogroup"] {
                display: flex;
                flex-direction: column;
            }
            .student-shell .stRadio label {
                width: 100%;
                min-height: 50px;
                font-size: 1rem;
            }
            .stButton > button, .stDownloadButton > button {
                width: 100%;
                min-height: 46px;
            }
            .teacher-link input {
                font-size: 0.82rem !important;
            }
            .quiz-meta-grid {
                grid-template-columns: 1fr 1fr;
            }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def get_query_params() -> dict:
    try:
        return dict(st.query_params)
    except Exception:
        return st.experimental_get_query_params()


def set_query_params(**kwargs) -> None:
    try:
        st.query_params.clear()
        for key, value in kwargs.items():
            st.query_params[key] = value
    except Exception:
        st.experimental_set_query_params(**kwargs)


def app_url_with_payload(token: str) -> str:
    try:
        base = st.context.url.split("?")[0]
    except Exception:
        base = ""
    return f"{base}?mode=student&a={token}" if base else f"?mode=student&a={token}"


def default_gas_url() -> str:
    params = get_query_params()
    saved = params.get("gas", "")
    if isinstance(saved, list):
        saved = saved[0] if saved else ""
    if saved:
        return saved
    try:
        return st.secrets.get("GAS_WEB_APP_URL", "")
    except Exception:
        return ""


def lesson_label_for_topic(topic: str) -> str:
    lesson_key = TOPIC_TO_LESSON.get(topic, "")
    return LESSONS.get(lesson_key, {}).get("label", "")


def remember_gas_url(gas_url: str) -> None:
    if not gas_url.strip():
        return
    try:
        st.query_params["gas"] = gas_url.strip()
    except Exception:
        st.experimental_set_query_params(gas=gas_url.strip())


def render_question(i: int, q: dict, key_prefix: str, allow_empty: bool = False):
    st.markdown(f"**Câu {i}.** {q['prompt']}")
    if q["qtype"] == "true_false":
        answers: dict[str, str | None] = {}
        for statement in q.get("statements") or []:
            st.markdown('<div class="tf-row">', unsafe_allow_html=True)
            st.markdown(f"**{statement['label']})** {statement['text']}")
            answers[statement["id"]] = st.radio(
                f"Chọn đúng/sai cho ý {statement['label']}",
                ["Đúng", "Sai"],
                key=f"{key_prefix}_{statement['id']}",
                horizontal=True,
                index=None if allow_empty else 0,
                label_visibility="collapsed",
            )
            st.markdown("</div>", unsafe_allow_html=True)
        return answers
    if q["qtype"] == "short_answer":
        st.markdown(
            '<div class="short-answer-note">Nhập đáp án bằng dấu phẩy thập phân. Nếu đáp án âm, chọn dấu “-”, rồi nhập 4 ký tự vào 4 ô.</div>',
            unsafe_allow_html=True,
        )
        sign_col, c1, c2, c3, c4 = st.columns([0.8, 1, 1, 1, 1])
        with sign_col:
            sign = st.selectbox("Dấu", ["", "-"], key=f"{key_prefix}_{q['id']}_sign", label_visibility="collapsed")
        chars = []
        for idx, col in enumerate([c1, c2, c3, c4], start=1):
            with col:
                value = st.text_input(
                    f"Ký tự {idx}",
                    key=f"{key_prefix}_{q['id']}_char_{idx}",
                    max_chars=1,
                    label_visibility="collapsed",
                    placeholder="",
                )
                chars.append(value)
        return sign + "".join(chars)
    return st.radio(
        "Chọn đáp án",
        q["options"],
        key=f"{key_prefix}_{q['id']}",
        horizontal=False,
        index=None if allow_empty else 0,
        label_visibility="collapsed",
    )


def submit_to_google_script(gas_url: str, payload: dict) -> tuple[bool, str]:
    if not gas_url:
        return False, "Chưa cấu hình Google Apps Script URL."
    try:
        response = requests.post(gas_url, json=payload, timeout=20)
        if 200 <= response.status_code < 300:
            return True, response.text[:300]
        return False, f"HTTP {response.status_code}: {response.text[:300]}"
    except Exception as exc:
        return False, str(exc)


def normalize_short_answer(value: str | None) -> str:
    if value is None:
        return ""
    normalized = str(value).strip().replace(" ", "").replace(".", ",")
    if normalized.startswith("+"):
        normalized = normalized[1:]
    return normalized


def answer_is_missing(q: dict, selected) -> bool:
    if q["qtype"] == "true_false":
        return any(value is None for value in (selected or {}).values())
    if q["qtype"] == "short_answer":
        return normalize_short_answer(selected).replace("-", "") == ""
    return selected is None


def score_answer(index: int, q: dict, selected) -> tuple[int, int, list[dict]]:
    if q["qtype"] == "true_false":
        score = 0
        details = []
        selected_map = selected or {}
        for statement in q.get("statements") or []:
            chosen = selected_map.get(statement["id"])
            ok = chosen == statement["answer"]
            score += int(ok)
            details.append(
                {
                    "index": f"{index}{statement['label']}",
                    "question": f"{q['prompt']} | {statement['label']}) {statement['text']}",
                    "selected": chosen,
                    "correct_answer": statement["answer"],
                    "is_correct": ok,
                    "explanation": statement["explanation"],
                }
            )
        return score, len(q.get("statements") or []), details

    if q["qtype"] == "short_answer":
        chosen = normalize_short_answer(selected)
        correct = normalize_short_answer(q["answer"])
        ok = chosen == correct
        return (
            int(ok),
            1,
            [
                {
                    "index": index,
                    "question": q["prompt"],
                    "selected": chosen,
                    "correct_answer": correct,
                    "is_correct": ok,
                    "explanation": q["explanation"],
                }
            ],
        )

    ok = selected == q["answer"]
    return (
        int(ok),
        1,
        [
            {
                "index": index,
                "question": q["prompt"],
                "selected": selected,
                "correct_answer": q["answer"],
                "is_correct": ok,
                "explanation": q["explanation"],
            }
        ],
    )


def plain_math_text(text: str) -> str:
    replacements = {
        "\\mathbb{R}": "R",
        "\\infty": "∞",
        "\\cup": "∪",
        "\\ge": "≥",
        "\\le": "≤",
        "\\Delta": "Δ",
        "\\sqrt": "√",
        "\\frac": "",
        "\\cdot": ".",
        "\\Leftrightarrow": "⇔",
    }
    cleaned = text.replace("$", "")
    for old, new in replacements.items():
        cleaned = cleaned.replace(old, new)
    cleaned = re.sub(r"\\[a-zA-Z]+", "", cleaned)
    cleaned = cleaned.replace("{", "").replace("}", "")
    return cleaned


def build_docx_bytes(payload: dict, questions: list) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(plain_math_text(payload["title"]).upper())
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 64, 175)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"{lesson_label_for_topic(payload['topic'])} | {TOPICS[payload['topic']]}").bold = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Mã đề: {payload['seed']} | Mã bài: {payload['assignment_id']} | Loại câu hỏi: {QUESTION_TYPES[payload['qtype']]}"
    )

    doc.add_paragraph("Họ và tên: ...............................................  Lớp: ................  Ngày: ....../....../......")

    for index, q in enumerate(questions, 1):
        q_para = doc.add_paragraph()
        q_para.paragraph_format.space_before = Pt(8)
        q_para.add_run(f"Câu {index}. ").bold = True
        q_para.add_run(plain_math_text(q.prompt))
        if q.qtype == "mcq":
            for label, option in zip(["A", "B", "C", "D"], q.options):
                doc.add_paragraph(f"{label}. {plain_math_text(option)}", style="List Bullet")
        elif q.qtype == "true_false":
            for statement in q.statements or []:
                doc.add_paragraph(f"{statement['label']}) {plain_math_text(statement['text'])}", style="List Bullet")
        else:
            doc.add_paragraph("Đáp án: □ □ □ □")

    doc.add_page_break()
    answer_title = doc.add_paragraph()
    answer_title.add_run("ĐÁP ÁN VÀ HƯỚNG DẪN").bold = True
    for index, q in enumerate(questions, 1):
        item = doc.add_paragraph()
        item.add_run(f"Câu {index}. ").bold = True
        item.add_run(f"Đáp án: {plain_math_text(q.answer)}")
        if q.statements:
            for statement in q.statements:
                doc.add_paragraph(
                    f"{statement['label']}) {statement['answer']} - {plain_math_text(statement['explanation'])}",
                    style="List Bullet",
                )
        else:
            doc.add_paragraph(plain_math_text(q.explanation))

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def google_sheet_help() -> None:
    with st.expander("Thiết lập Google Sheet nhận kết quả"):
        st.markdown(
            """
            1. Tạo một Google Sheet mới.
            2. Vào **Extensions -> Apps Script**.
            3. Dán nội dung file `google_apps_script.gs` trong repo GitHub.
            4. Bấm **Deploy -> New deployment -> Web app**.
            5. Chọn **Execute as: Me** và **Who has access: Anyone**.
            6. Copy **Web app URL** rồi dán vào ô cấu hình trong app.

            Sau khi học sinh nộp bài, script sẽ ghi vào sheet `KetQua` với các cột:
            `Thứ tự`, `Họ tên học sinh`, `Lớp`, `Ngày làm bài`, `Kết quả`, `Kết quả từng câu`.
            """
        )


def teacher_page() -> None:
    st.markdown(
        """
        <div class="app-hero">
            <h1>Tạo bài tập online Chương 6 - Toán 10</h1>
            <p>Chọn dạng toán, loại câu hỏi và số câu; app sinh đề, tạo link làm bài trên điện thoại và gửi kết quả về Google Sheet.</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    if "generated_count" not in st.session_state:
        st.session_state.generated_count = 0

    tab_setup, tab_preview, tab_stats = st.tabs(["Thiết lập & chia sẻ", "Xem trước đề", "Thống kê"])

    with tab_setup:
        with st.form("teacher_settings_form"):
            left, right = st.columns([1.1, 0.9], gap="large")
            with left:
                st.subheader("Thông tin đề")
                title = st.text_input("Tên bài tập", "Luyện tập Chương 6 - Toán 10")
                teacher = st.text_input("Tên giáo viên", "")
                lesson_keys = list(LESSONS)
                lesson = st.selectbox("Chọn bài", lesson_keys, format_func=lambda k: LESSONS[k]["label"])
                topic_options = LESSONS[lesson]["topics"]
                topic = st.selectbox("Dạng cụ thể", topic_options, format_func=lambda k: TOPICS[k])
            with right:
                st.subheader("Cấu hình sinh đề")
                qtype = st.radio("Loại câu hỏi", list(QUESTION_TYPES), format_func=lambda k: QUESTION_TYPES[k])
                count = st.slider("Số câu", min_value=3, max_value=30, value=10)
                seed = st.number_input("Mã đề/seed", min_value=1, max_value=999999, value=12345, step=1)
                gas_url = st.text_input(
                    "Google Apps Script Web App URL",
                    value=default_gas_url(),
                    placeholder="https://script.google.com/macros/s/.../exec",
                )
            generate = st.form_submit_button("Tạo / cập nhật bài tập", type="primary", use_container_width=True)

        if generate or "teacher_payload" not in st.session_state:
            if generate:
                remember_gas_url(gas_url)
                st.session_state.generated_count += 1
                st.toast("Đã tạo bài tập và lưu URL Google Sheet cho trang này.")
            st.session_state.teacher_payload = assignment_payload(topic, qtype, count, title, teacher, gas_url, seed)

        payload = st.session_state.teacher_payload
        questions = generate_questions(payload["topic"], payload["qtype"], payload["count"], payload["seed"])
        token = encode_payload(payload)
        student_link = app_url_with_payload(token)
        lesson_label = lesson_label_for_topic(payload["topic"])

        st.markdown(
            f"""
            <div class="quiz-meta-grid">
                <div class="quiz-meta-item"><span>Bài học</span><strong>{lesson_label or "Chương 6"}</strong></div>
                <div class="quiz-meta-item"><span>Dạng cụ thể</span><strong>{TOPICS[payload['topic']]}</strong></div>
                <div class="quiz-meta-item"><span>Số câu</span><strong>{payload['count']}</strong></div>
                <div class="quiz-meta-item"><span>Mã đề</span><strong>{payload['seed']}</strong></div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        share_col, status_col = st.columns([1.35, 0.65], gap="large")
        with share_col:
            st.markdown('<div class="color-card">', unsafe_allow_html=True)
            st.subheader("Link giao bài")
            st.markdown('<div class="teacher-link">', unsafe_allow_html=True)
            st.text_input("Link học sinh", student_link)
            st.markdown("</div>", unsafe_allow_html=True)
            json_data = json.dumps(
                {"assignment": payload, "questions": questions_to_dicts(questions)},
                ensure_ascii=False,
                indent=2,
            )
            docx_data = None
            docx_error = ""
            try:
                docx_data = build_docx_bytes(payload, questions)
            except ModuleNotFoundError:
                docx_error = "Chưa cài thư viện python-docx. Khi deploy, Streamlit sẽ cài theo requirements.txt."
            d1, d2 = st.columns(2)
            with d1:
                st.download_button(
                    "Tải JSON",
                    data=json_data,
                    file_name=f"bai-tap-chuong-6-{payload['assignment_id']}.json",
                    mime="application/json",
                    use_container_width=True,
                )
            with d2:
                if docx_data:
                    st.download_button(
                        "Tải đề Word (.docx)",
                        data=docx_data,
                        file_name=f"de-chuong-6-{payload['assignment_id']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )
                else:
                    st.warning(docx_error or "Chưa tạo được file Word.")
            st.markdown("</div>", unsafe_allow_html=True)
        with status_col:
            with st.expander("Hướng dẫn & trạng thái", expanded=True):
                if payload["gas_url"]:
                    st.success("Đã có Google Apps Script URL.")
                else:
                    st.warning("Chưa có Google Apps Script URL.")
                st.info("Khi deploy trên Streamlit Community Cloud, link giao bài sẽ là link thật của app.")
            google_sheet_help()

    payload = st.session_state.teacher_payload
    questions = generate_questions(payload["topic"], payload["qtype"], payload["count"], payload["seed"])

    with tab_preview:
        st.subheader("Xem trước đề")
        for i, q in enumerate(questions, 1):
            st.markdown('<div class="question-card">', unsafe_allow_html=True)
            st.markdown(f"**Câu {i}.** {q.prompt}")
            if q.qtype == "mcq":
                for label, option in zip(["A", "B", "C", "D"], q.options):
                    st.markdown(f"{label}. {option}")
            elif q.qtype == "true_false":
                for statement in q.statements or []:
                    st.markdown(f"**{statement['label']})** {statement['text']}")
            else:
                st.markdown("Học sinh nhập đáp số vào 4 ô ký tự, dùng dấu phẩy thập phân nếu có.")
            with st.expander("Đáp án và hướng dẫn"):
                st.markdown(f"**Đáp án:** {q.answer}")
                st.markdown(q.explanation)
            st.markdown("</div>", unsafe_allow_html=True)

    with tab_stats:
        st.subheader("Thống kê sử dụng")
        m1, m2, m3 = st.columns(3)
        m1.metric("Lượt sinh đề trong phiên", st.session_state.generated_count)
        m2.metric("Số câu của đề hiện tại", payload["count"])
        m3.metric("Mã bài hiện tại", payload["assignment_id"])
        st.info("Phổ điểm và tỷ lệ đúng/sai từng câu có thể bổ sung ở bước sau khi có cơ chế đọc dữ liệu ngược từ Google Sheet hoặc file CSV xuất từ Sheet.")


def student_page(payload: dict) -> None:
    st.markdown('<div class="student-shell">', unsafe_allow_html=True)
    st.markdown(
        f"""
        <div class="app-hero">
            <h1>{payload["title"]}</h1>
            <p>Mã bài: <b>{payload['assignment_id']}</b> &nbsp;|&nbsp; Bài: <b>{lesson_label_for_topic(payload['topic']) or 'Chương 6'}</b> &nbsp;|&nbsp; Dạng: <b>{TOPICS[payload['topic']]}</b> &nbsp;|&nbsp; Số câu: <b>{payload['count']}</b></p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    questions = questions_to_dicts(generate_questions(payload["topic"], payload["qtype"], payload["count"], payload["seed"]))

    with st.form("student_quiz_form"):
        st.subheader("Thông tin học sinh")
        col1, col2 = st.columns(2)
        with col1:
            student_name = st.text_input("Họ và tên học sinh")
        with col2:
            student_class = st.text_input("Lớp")

        st.subheader("Bài làm")
        answers: dict[str, str | None] = {}
        for i, q in enumerate(questions, 1):
            with st.container(border=True):
                answers[q["id"]] = render_question(i, q, "student", allow_empty=True)

        submitted = st.form_submit_button("Nộp bài", type="primary", use_container_width=True)

    if submitted:
        if not student_name.strip() or not student_class.strip():
            st.error("Em cần nhập đầy đủ họ tên và lớp trước khi nộp bài.")
            st.stop()
        unanswered = [i for i, q in enumerate(questions, 1) if answer_is_missing(q, answers.get(q["id"]))]
        if unanswered:
            st.error("Em chưa hoàn thành câu: " + ", ".join(map(str, unanswered)))
            st.stop()

        score = 0
        total_units = 0
        details = []
        for i, q in enumerate(questions, 1):
            selected = answers.get(q["id"])
            item_score, item_total, item_details = score_answer(i, q, selected)
            score += item_score
            total_units += item_total
            details.extend(item_details)

        result = {
            "submitted_at": datetime.now().isoformat(timespec="seconds"),
            "assignment_id": payload["assignment_id"],
            "assignment_title": payload["title"],
            "teacher": payload.get("teacher", ""),
            "topic": TOPICS[payload["topic"]],
            "question_type": QUESTION_TYPES[payload["qtype"]],
            "seed": payload["seed"],
            "student_name": student_name.strip(),
            "student_class": student_class.strip(),
            "score": score,
            "total": total_units,
            "percent": round(score * 100 / total_units, 2),
            "details": details,
        }

        st.success(f"Em đạt {score}/{total_units} ý đúng ({result['percent']}%).")
        ok, message = submit_to_google_script(payload.get("gas_url", ""), result)
        if ok:
            st.success("Đã gửi kết quả về Google Sheet.")
        else:
            st.warning(f"Chưa gửi được về Google Sheet: {message}")

        with st.expander("Xem đáp án và lời giải"):
            for item in details:
                icon = "Đúng" if item["is_correct"] else "Sai"
                st.markdown(f"**Câu {item['index']} - {icon}**")
                st.markdown(f"Em chọn: `{item['selected']}` | Đáp án đúng: `{item['correct_answer']}`")
                st.markdown(item["explanation"])

    st.markdown("</div>", unsafe_allow_html=True)


def main() -> None:
    apply_mobile_css()
    params = get_query_params()
    mode = params.get("mode", "teacher")
    token = params.get("a", "")
    if isinstance(mode, list):
        mode = mode[0]
    if isinstance(token, list):
        token = token[0]

    if mode == "student" and token:
        try:
            student_page(decode_payload(token))
        except Exception as exc:
            st.error(f"Link bài tập không hợp lệ: {exc}")
            if st.button("Về trang giáo viên"):
                set_query_params(mode="teacher")
    else:
        teacher_page()


if __name__ == "__main__":
    main()
